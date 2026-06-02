from docx import Document
from docx.shared import Inches
from io import BytesIO

def create_word_document(summary_text):
    """
    Exportador Word: Genera fichas de estudio con márgenes estrechos,
    traduce flechas lógicas y optimiza la densidad visual del texto.
    """
    doc = Document()
    
    # Optimización de página: Márgenes estrechos
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Título limpio sin el término rechazado
    doc.add_heading('Fichas de Estudio Clínico', 0)
    
    lines = summary_text.split('\n')
    in_table = False
    table = None
    
    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue
            
        # Reemplazo tipográfico de flechas lógicas
        line_str = line_str.replace('-->', '→').replace('->', '→')
            
        # =====================================================================
        # DETECCIÓN Y PROCESAMIENTO DE MATRICES (TABLAS)
        # =====================================================================
        if line_str.startswith('|') and line_str.endswith('|'):
            if '---' in line_str:
                continue 
                
            celdas = [c.strip() for c in line_str.split('|')[1:-1]]
            
            if not in_table:
                in_table = True
                table = doc.add_table(rows=0, cols=len(celdas))
                table.style = 'Table Grid'
                
                row = table.add_row()
                for i, cell_text in enumerate(celdas):
                    clean_text = cell_text.replace('**', '').replace('*', '').strip()
                    cell = row.cells[i]
                    p = cell.paragraphs[0]
                    run = p.add_run(clean_text)
                    run.bold = True 
            else:
                row = table.add_row()
                for i, cell_text in enumerate(celdas):
                    cell = row.cells[i]
                    p = cell.paragraphs[0]
                    
                    partes = cell_text.replace('<br>', '\n').split('\n')
                    
                    for idx, parte in enumerate(partes):
                        parte_limpia = parte.strip()
                        if not parte_limpia:
                            continue
                            
                        if idx > 0 or p.text:
                            p = cell.add_paragraph()
                            
                        if parte_limpia.startswith('- ') or parte_limpia.startswith('* '):
                            p.style = 'List Bullet'
                            texto_final = parte_limpia[2:].strip()
                        else:
                            texto_final = parte_limpia
                            
                        texto_final_limpio = texto_final.replace('**', '').replace('*', '')
                        run = p.add_run(texto_final_limpio)
                        
                        if i == 0:
                            run.bold = True
            continue
        else:
            in_table = False 
            
        # =====================================================================
        # PROCESAMIENTO DE TÍTULOS Y TEXTO LIBRE
        # =====================================================================
        if line_str.startswith('### '):
            doc.add_heading(line_str.replace('### ', '').replace('**', '').strip(), level=2)
        elif line_str.startswith('## '):
            doc.add_heading(line_str.replace('## ', '').replace('**', '').strip(), level=1)
        elif line_str.startswith('# '):
            doc.add_heading(line_str.replace('# ', '').replace('**', '').strip(), level=0)
        else:
            p_text = line_str.replace('**', '').replace('*', '').strip()
            doc.add_paragraph(p_text)
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
